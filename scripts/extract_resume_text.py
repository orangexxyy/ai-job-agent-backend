"""从本地简历文件提取纯文本，生成 current_resume.txt。

支持 docx、文本型 PDF、txt、md。脚本只做本地文件读取和文本落盘，
不调用 LLM，不写数据库，也不修改 candidate_profile。
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import NamedTuple


DEFAULT_OUTPUT = Path("docs/input/current_resume.txt")
DEFAULT_REPORT = Path("docs/input/resume_extract_report.md")
SUPPORTED_SUFFIXES = {".docx", ".pdf", ".txt", ".md"}
SCANNED_PDF_TEXT_THRESHOLD = 80


class ExtractResult(NamedTuple):
    text: str
    file_type: str
    pdf_pages: int | None = None
    suspected_scanned_pdf: bool = False
    read_encoding: str | None = None


def _clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip() + "\n" if text.strip() else ""


def _read_text_file(path: Path) -> ExtractResult:
    for encoding in ("utf-8", "gbk"):
        try:
            return ExtractResult(
                text=path.read_text(encoding=encoding),
                file_type=path.suffix.lower().lstrip("."),
                read_encoding=encoding,
            )
        except UnicodeDecodeError:
            continue
    raise ValueError("无法按 utf-8 或 gbk 读取文本文件")


def _read_docx(path: Path) -> ExtractResult:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少依赖 python-docx，请先安装 requirements.txt") from exc

    document = Document(path)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            lines.append(value)

    for table in document.tables:
        for row in table.rows:
            row_values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text and cell.text.strip()
            ]
            if row_values:
                lines.append(" | ".join(row_values))

    return ExtractResult(text="\n".join(lines), file_type="docx")


def _read_pdf(path: Path) -> ExtractResult:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("缺少依赖 pypdf，请先安装 requirements.txt") from exc

    reader = PdfReader(str(path))
    page_texts: list[str] = []

    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        page_text = page_text.strip()
        if page_text:
            page_texts.append(f"--- Page {index} ---\n{page_text}")

    text = "\n\n".join(page_texts)
    suspected_scanned_pdf = len(text.strip()) < SCANNED_PDF_TEXT_THRESHOLD
    return ExtractResult(
        text=text,
        file_type="pdf",
        pdf_pages=len(reader.pages),
        suspected_scanned_pdf=suspected_scanned_pdf,
    )


def extract_resume_text(input_path: Path) -> ExtractResult:
    suffix = input_path.suffix.lower()

    if suffix not in SUPPORTED_SUFFIXES:
        supported = " / ".join(sorted(SUPPORTED_SUFFIXES))
        raise ValueError(f"不支持的文件类型：{suffix or '无后缀'}。支持 {supported}")

    if suffix == ".docx":
        return _read_docx(input_path)
    if suffix == ".pdf":
        return _read_pdf(input_path)
    return _read_text_file(input_path)


def _write_report(input_path: Path, output_path: Path, result: ExtractResult) -> None:
    report_path = DEFAULT_REPORT
    report_path.parent.mkdir(parents=True, exist_ok=True)

    scan_note = "是" if result.suspected_scanned_pdf else "否"
    pdf_pages = str(result.pdf_pages) if result.pdf_pages is not None else "不适用"

    suggestions = [
        "下一步可以人工检查 `docs/input/current_resume.txt` 是否完整。",
        "确认无误后，可基于该文件整理新版简历和 candidate_profile 草稿。",
    ]
    if result.suspected_scanned_pdf:
        suggestions.insert(
            0,
            "PDF 提取文本过短，可能是扫描版 PDF；当前不支持 OCR，请使用文本型 PDF 或 docx。",
        )

    content = "\n".join(
        [
            "# Resume Extract Report",
            "",
            f"- 输入文件路径：`{input_path}`",
            f"- 文件类型：`{result.file_type}`",
            f"- 输出文件路径：`{output_path}`",
            f"- 提取字符数：`{len(result.text.strip())}`",
            f"- PDF 页数：`{pdf_pages}`",
            f"- 是否疑似扫描版 PDF：`{scan_note}`",
            f"- 文本读取编码：`{result.read_encoding or '不适用'}`",
            "",
            "## 下一步建议",
            "",
            *[f"- {item}" for item in suggestions],
            "",
        ]
    )
    report_path.write_text(content, encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="从 docx / 文本型 PDF / txt / md 简历中提取文本。"
    )
    parser.add_argument("input_file", help="输入简历文件路径")
    parser.add_argument(
        "--output",
        default=str(DEFAULT_OUTPUT),
        help="输出 current_resume.txt 路径，默认 docs/input/current_resume.txt",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path = Path(args.input_file)
    output_path = Path(args.output)

    Path("docs/input").mkdir(parents=True, exist_ok=True)
    Path("docs/input/resume_source").mkdir(parents=True, exist_ok=True)

    if not input_path.exists():
        print(f"错误：输入文件不存在：{input_path}", file=sys.stderr)
        return 1
    if not input_path.is_file():
        print(f"错误：输入路径不是文件：{input_path}", file=sys.stderr)
        return 1

    try:
        result = extract_resume_text(input_path)
        cleaned_text = _clean_text(result.text)
        normalized_result = result._replace(text=cleaned_text)
    except (ValueError, RuntimeError) as exc:
        print(f"错误：{exc}", file=sys.stderr)
        return 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(cleaned_text, encoding="utf-8")
    _write_report(input_path, output_path, normalized_result)

    print(f"已生成：{output_path}")
    print(f"已生成：{DEFAULT_REPORT}")
    print(f"提取字符数：{len(cleaned_text.strip())}")
    if normalized_result.suspected_scanned_pdf:
        print("提示：PDF 提取文本过短，可能是扫描版；当前不支持 OCR。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
